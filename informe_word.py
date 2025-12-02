from io import BytesIO
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt


# ============================
# UTILIDADES DE FORMATO
# ============================

def agregar_titulo(doc: Document, texto: str, size: int = 16, bold: bool = True):
    """
    Agrega un título o subtítulo al documento, con tamaño y negrita configurables.
    """
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = bold
    r.font.size = Pt(size)


def agregar_parrafo(doc: Document, texto: str, size: int = 11, bold: bool = False):
    """
    Agrega un párrafo estándar al documento.
    """
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = bold
    r.font.size = Pt(size)


def agregar_tabla_criterios(doc: Document, criterios: Dict[str, Any]):
    """
    Agrega una tabla con los criterios (C1, C2, ..., C13) y sus puntajes.
    """
    if not criterios:
        agregar_parrafo(doc, "No se encontraron criterios evaluados.")
        return

    # Cabecera de tabla
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Criterio"
    hdr_cells[1].text = "Puntaje"

    # Ordenamos los criterios por nombre (C1, C2, ..., C13) para que quede prolijo
    for k in sorted(criterios.keys()):
        v = criterios[k]
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)


def agregar_incongruencias(doc: Document, incong: Any):
    """
    Agrega al informe la sección de incongruencias detectadas.

    - Si 'incong' es:
      * None o lista vacía → indica que no se detectaron.
      * List[dict] → formato estructurado del módulo incongruencias.py
      * str / dict genérico → se imprime tal cual, como respaldo.
    """
    if not incong:
        agregar_parrafo(doc, "No se registraron incongruencias detectadas.")
        return

    # Caso 1: lista estructurada proveniente de analizar_incongruencias()
    if isinstance(incong, list) and incong and isinstance(incong[0], dict):
        for i, item in enumerate(incong, 1):
            tipo = item.get("tipo", "Incongruencia sin tipo especificado")
            parrafos = item.get("parrafos", [])
            detalle = item.get("detalle", "")
            extractos: List[str] = item.get("extractos", [])

            # Encabezado numerado
            agregar_parrafo(doc, f"{i}. {tipo}", bold=True)

            # Párrafos afectados
            if parrafos:
                parrafos_str = ", ".join(str(n) for n in parrafos)
                agregar_parrafo(doc, f"Párrafos involucrados: {parrafos_str}")

            # Detalle
            if detalle:
                agregar_parrafo(doc, f"Detalle: {detalle}")

            # Extractos
            if extractos:
                agregar_parrafo(doc, "Extractos relevantes:", bold=True)
                for ex in extractos:
                    p = doc.add_paragraph(style=None)
                    run = p.add_run(f"- {ex}")
                    run.font.size = Pt(10)

            # Línea en blanco entre incongruencias
            agregar_parrafo(doc, "")
        return

    # Caso 2: si solo es un string (respaldo)
    if isinstance(incong, str):
        agregar_parrafo(doc, incong)
        return

    # Caso 3: si es un dict genérico
    if isinstance(incong, dict):
        for k, v in incong.items():
            agregar_parrafo(doc, f"- {k}: {v}")
        return

    # Caso 4: cualquier otro tipo → lo imprimimos en bruto
    agregar_parrafo(doc, str(incong))


# ============================
# FUNCIÓN PRINCIPAL
# ============================

def generar_informe(texto: str, resultados: Dict[str, Any], incong: Any) -> bytes:
    """
    Genera un informe en formato .docx (devuelto como bytes) a partir de:

    - texto: sentencia analizada (por ahora solo se usa para contexto futuro).
    - resultados: dict devuelto por el evaluador (criterios, ICI, interpretación).
    - incong: lista de incongruencias devuelta por analizar_incongruencias().

    Retorna: bytes listos para usar en st.download_button en Streamlit.
    """
    doc = Document()

    # ============================
    # PORTADA
    # ============================
    agregar_titulo(doc, "INFORME DE COHERENCIA INDICIARIA – ICI V7", size=18)
    agregar_parrafo(doc, "Sistema de Auditoría Indiciaria ACRJ – Versión 7.")
    agregar_parrafo(doc, "")
    agregar_parrafo(
        doc,
        "Este informe resume el análisis automatizado realizado sobre la resolución o sentencia "
        "cargada, aplicando criterios de coherencia indiciaria y reglas de detección de "
        "incongruencias lógicas y normativas.",
    )
    doc.add_page_break()

    # ============================
    # 1. RESUMEN ICI
    # ============================
    agregar_titulo(doc, "1. RESUMEN DEL ÍNDICE DE COHERENCIA INDICIARIA", size=14)

    # Extraer datos del dict de resultados
    criterios = resultados.get("criterios", {}) if isinstance(resultados, dict) else {}
    ici_sin = resultados.get("ICI_sin_penalizacion", None) if isinstance(resultados, dict) else None
    ici_aj = resultados.get("ICI_ajustado", None) if isinstance(resultados, dict) else None
    interpretacion = resultados.get("interpretacion", "") if isinstance(resultados, dict) else ""

    if ici_sin is not None:
        agregar_parrafo(doc, f"ICI sin penalización: {ici_sin}", bold=True)
    if ici_aj is not None:
        agregar_parrafo(doc, f"ICI ajustado (con frenos de emergencia): {ici_aj}", bold=True)

    agregar_parrafo(doc, "")
    agregar_parrafo(doc, "Interpretación cualitativa:", bold=True)
    if interpretacion:
        agregar_parrafo(doc, interpretacion)
    else:
        agregar_parrafo(doc, "No se ha generado una interpretación cualitativa.")

    # ============================
    # 2. DETALLE DE CRITERIOS C1 – C13
    # ============================
    doc.add_page_break()
    agregar_titulo(doc, "2. DETALLE DE CRITERIOS C1 – C13", size=14)
    agregar_parrafo(
        doc,
        "Se detallan los puntajes asignados a cada criterio de coherencia indiciaria "
        "(C1, C2, ..., C13). Cuando alguno no aparezca, se debe a que el sistema no "
        "pudo detectarlo en el texto analizado.",
    )
    agregar_parrafo(doc, "")
    agregar_tabla_criterios(doc, criterios)

    # ============================
    # 3. INCONGRUENCIAS DETECTADAS
    # ============================
    doc.add_page_break()
    agregar_titulo(doc, "3. INCONGRUENCIAS DETECTADAS", size=14)
    agregar_parrafo(
        doc,
        "Se listan las principales incongruencias lógicas, tensiones probatorias o "
        "saltos argumentativos detectados por el módulo de análisis de incongruencias "
        "(Reglas 0–9). Cada ítem incluye los párrafos afectados y extractos relevantes.",
    )
    agregar_parrafo(doc, "")
    agregar_incongruencias(doc, incong)

    # ============================
    # 4. NOTAS METODOLÓGICAS
    # ============================
    doc.add_page_break()
    agregar_titulo(doc, "4. NOTAS METODOLÓGICAS", size=14)
    agregar_parrafo(
        doc,
        "El sistema ACRJ–ICI V7 evalúa la coherencia indiciaria de las resoluciones "
        "mediante reglas heurísticas y patrones lingüísticos inspirados en el método "
        "indiciario (pluralidad de indicios, fiabilidad, nexo lógico, hipótesis "
        "alternativas, coherencia interna y externa, ausencia de circularidad, etc.).",
    )
    agregar_parrafo(
        doc,
        "El resultado NO sustituye el juicio crítico del abogado defensor ni del tribunal, "
        "sino que ofrece un mapa de riesgos argumentativos para orientar la revisión humana. "
        "Debe interpretarse siempre como una herramienta de apoyo, no como una decisión "
        "automática sobre culpabilidad o inocencia.",
    )
    agregar_parrafo(
        doc,
        "Se recomienda revisar especialmente los criterios con puntajes bajos (por debajo de 60), "
        "así como las incongruencias lógicas, saltos probatorios y contradicciones internas "
        "identificadas por el sistema, a fin de fundamentar recursos de nulidad, apelación o "
        "acciones extraordinarias (revisión, habeas corpus, etc.) cuando corresponda.",
    )

    # GENERAR BYTES DEL DOCUMENTO PARA STREAMLIT
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
